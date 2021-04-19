from docx import Document as DocumentXML
from docx.enum.text import WD_BREAK

from workplace.common.utils import get_report_full_path, get_string_dates


def fill_title(docx_file, string_date):
    paragraph = docx_file.add_paragraph()
    paragraph.alignment = 1
    paragraph.add_run('Итоговые часы за %s' % string_date).bold = True
    paragraph.add_run().add_break(WD_BREAK.LINE)


def get_row_cells(table, row_number):
    return table.rows[row_number].cells


def merge_cells(table, start_column, end_column, start_row, end_row):
    start_cell = table.cell(start_row, start_column)
    end_cell = table.cell(end_row, end_column)
    start_cell.merge(end_cell)


def create_docx_month_report(items, start, end, user_id, report_id):
    string_date = get_string_dates(start, end)
    full_path = get_report_full_path(user_id, report_id, string_date, 'docx')
    docx_file = DocumentXML()
    fill_title(docx_file, string_date)
    table = docx_file.add_table(rows=1, cols=3)
    row_number = 0
    row_cells = get_row_cells(table, row_number)
    row_cells[0].text = 'Проект'
    row_cells[1].text = 'Направление'
    row_cells[2].text = 'Время'
    table.add_row()
    row_number += 1
    row_cells = get_row_cells(table, row_number)
    total_duration = 0
    for item in items:
        start_row = row_number
        row_cells[0].text = item.get('project')
        for direction in item.get('directions'):
            duration = direction.get('duration') / 60
            total_duration += direction.get('duration')
            row_cells[1].text = direction.get('direction')
            row_cells[2].text = '%s' % round(duration, 2)
            table.add_row()
            row_number += 1
            row_cells = get_row_cells(table, row_number)

        end_row = row_number - 1
        if end_row > start_row:
            merge_cells(table, 0, 0, start_row, end_row)
    row_cells = get_row_cells(table, row_number)
    row_cells[0].text = 'Итого'
    row_cells[2].text = '%s' % round(total_duration / 60, 2)
    merge_cells(table, 0, 1, row_number, row_number)
    docx_file.save(full_path)
    return full_path
